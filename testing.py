from bs4 import BeautifulSoup
from numpy import concatenate
import requests
# from fpdf import FPDF
from io import StringIO
from html.parser import HTMLParser
# import urllib.request
import docx

class MLStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = StringIO()
    def handle_data(self, d):
        self.text.write(d)
    def get_data(self):
        return self.text.getvalue()

def strip_tags(html):
    s = MLStripper()
    s.feed(html)
    return s.get_data()

def crawl_content(post_url):
    html_text = requests.get(post_url).text
    soup = BeautifulSoup(html_text, 'html.parser')
    page_title = soup.title.string
    web_content = str(soup.find_all('div', attrs={'class':'nh-read__content post-body', 'id': 'js-read__content'})[0]).replace("<br/>","\n")
    content = strip_tags(web_content)
    return page_title, content


if __name__ == "__main__":
    # pdf = FPDF()
    # pdf.add_page()
    # pdf.add_font('TimeNewRoman', '', 'times.ttf', uni=True)
    # pdf.set_font('TimeNewRoman','',15)
    doc = docx.Document()
    data = ""
    for x in range(1051, 1141):
        url = "https://vtruyen.com/truyen/toan-chuc-phap-su-di-ban/chuong-{}".format(x)
        title, content = crawl_content(url)
        new_content = content.encode('utf-8', 'ignore').decode('utf-8').replace("— QUẢNG CÁO —","")
        chapter_title = title.replace("Đọc Toàn Chức Pháp Sư Dị Bản - ", "").replace(" online", "")
        # pdf.cell(200, 10, txt = chapter_title, ln = 1, align = 'C')
        # data = data + "\n" + new_content
        doc.add_heading(chapter_title, 3)
        doc.add_paragraph(new_content)
        doc.add_page_break()
        if x == 1140:
            start_chap = 1051
            end_chap = 1140

            # with open("F:\\Desktop\\ManyThings\\TCPS\\docs\\Chapter {} - {}.txt".format(start_chap, end_chap), "w", encoding="utf-8") as text_file:
            #     text_file.write(chapter_title)
            #     text_file.write(data)
            # data = ""
            doc.save('F:\\Desktop\\ManyThings\\TCPS\\Chapter {} - {}.docx'.format(start_chap, end_chap))
        # for txt in new_content.split('\n'):
        #     pdf.write(8, txt)
        #     pdf.ln(8)
        # pdf.add_page()

        # if x % 50 == 0:
        #     start_chap = x - 49
        #     end_chap = x
        #     pdf.output("TCPS\\new\{}.pdf".format("Chapter {0} - {1}".format(start_chap, end_chap)), 'F')
        #     pdf = FPDF()
        #     pdf.add_font('TimeNewRoman', '', 'times.ttf', uni=True)
        #     pdf.set_font('TimeNewRoman','',15)
        #     pdf.add_page()
